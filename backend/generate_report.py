import json
import os
from docx import Document
from docx.shared import Inches

# Create reports directory if it doesn't exist
if not os.path.exists('reports'):
    os.makedirs('reports')

# Load click data from JSON file
with open('clickData.json', 'r') as file:
    click_data = json.load(file)

# Create a new Word Document
doc = Document()
doc.add_heading('Click Event Report', 0)

# Iterate over the click data
for click in click_data:
    doc.add_heading(f"Click at ({click['x']}, {click['y']})", level=1)
    doc.add_paragraph(f"Timestamp: {click['timestamp']}")

    # Check if the screenshot path exists before adding it to the document
    if os.path.exists(click['screenshotPath']):
        doc.add_picture(click['screenshotPath'], width=Inches(5.0))
    else:
        doc.add_paragraph("Screenshot not found.")

# Save the document
doc.save('reports/ClickEventReport.docx')
print("Report generated successfully.")
